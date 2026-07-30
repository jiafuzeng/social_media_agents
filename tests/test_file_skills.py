from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from integrated_agent.gateway import GatewayAttachment, GatewayRequest
from integrated_agent.runtimes.agent import (
    AgentlyAgentRuntime,
    WorkspaceFileService,
)
from integrated_agent.storage import ArtifactStore


def test_xlsx_skill_returns_markdown_text(tmp_path: Path) -> None:
    source = tmp_path / "经营数据.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    assert worksheet is not None
    worksheet.title = "月度数据"
    worksheet.append(["月份", "净营收"])
    worksheet.append(["2025-01", 120])
    workbook.save(source)

    result = WorkspaceFileService(tmp_path / "workspace").process(source)

    assert result.operation_key == "xlsx_to_markdown"
    assert "工作表：月度数据" in result.text
    assert "| 月份 | 净营收 |" in result.text
    assert "| 2025-01 | 120 |" in result.text


def test_docx_skill_returns_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "复盘.docx"
    document = Document()
    document.add_heading("项目复盘", level=1)
    document.add_paragraph("本月重点是提升交付稳定性。")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "事项"
    table.rows[0].cells[1].text = "状态"
    table.rows[1].cells[0].text = "SSE接入"
    table.rows[1].cells[1].text = "完成"
    document.save(str(source))

    result = WorkspaceFileService(tmp_path / "workspace").process(source)

    assert result.operation_key == "docx_to_markdown"
    assert "# 项目复盘" in result.text
    assert "本月重点是提升交付稳定性。" in result.text
    assert "| 事项 | 状态 |" in result.text
    assert "| SSE接入 | 完成 |" in result.text


def test_markdown_skill_generates_pdf_and_streams_text(tmp_path: Path) -> None:
    source = tmp_path / "周报.md"
    source.write_text(
        "# 本周进展\n\n- 完成问数服务接入\n- 补齐接口压测\n",
        encoding="utf-8",
    )
    service = WorkspaceFileService(tmp_path / "workspace")
    runtime = AgentlyAgentRuntime(
        artifact_store=ArtifactStore(tmp_path / "artifacts"),
        public_base_url="http://127.0.0.1:8000",
        skills_root=Path(__file__).parents[1] / "skills",
        registry_root=tmp_path / "skills_registry",
        workspace_root=tmp_path / "workspace",
        file_service=service,
    )

    async def collect_events() -> list[tuple[str, dict]]:
        return [
            (event.type, event.data)
            async for event in runtime.stream(
                GatewayRequest(
                    "处理上传文件",
                    "course-user",
                    attachments=(
                        GatewayAttachment(source, source.name),
                    ),
                )
            )
        ]

    events = asyncio.run(collect_events())
    result = service.process(source)

    assert result.operation_key == "markdown_to_pdf"
    assert result.artifact_path is not None
    assert result.artifact_path.read_bytes().startswith(b"%PDF-")
    assert "处理后的文本预览" in result.text
    assert [event_type for event_type, _ in events] == [
        "run.created",
        "status.update",
        "message.delta",
        "artifact.ready",
        "run.completed",
    ]
    assert "本周进展" in str(events[2][1]["delta"])
    assert events[3][1]["artifact_id"]
    assert Path(str(events[3][1]["path"])).is_file()
