from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable
from xml.sax.saxutils import escape

from docx import Document
from openpyxl import load_workbook
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer


MAX_UPLOAD_BYTES = 20 * 1024 * 1024
MAX_TEXT_CHARS = 12_000


@dataclass(frozen=True)
class FileOperationResult:
    operation_key: str
    source_name: str
    text: str
    artifact_path: Path | None = None
    artifact_mime_type: str | None = None


FileHandler = Callable[[Path], FileOperationResult]


def _clean_component(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z._-]+", "_", value).strip("._")
    return cleaned[:100] or fallback


def _cell_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("|", "\\|").replace("\n", " ")


def _markdown_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    return [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
        *[
            "| " + " | ".join(row) + " |"
            for row in normalized[1:]
        ],
    ]


class WorkspaceFileService:
    def __init__(self, workspace_root: Path) -> None:
        self.workspace_root = workspace_root
        self.uploads_root = workspace_root / "uploads"
        self.artifacts_root = workspace_root / "artifacts"
        self._handlers: dict[str, FileHandler] = {
            ".xlsx": self._xlsx_to_markdown,
            ".docx": self._docx_to_markdown,
            ".md": self._markdown_to_pdf,
            ".markdown": self._markdown_to_pdf,
        }

    @property
    def supported_suffixes(self) -> tuple[str, ...]:
        return tuple(sorted(self._handlers))

    def save_upload(
        self,
        *,
        session_id: str,
        filename: str,
        content: bytes,
    ) -> Path:
        if not content:
            raise ValueError("上传文件为空")
        if len(content) > MAX_UPLOAD_BYTES:
            raise ValueError("上传文件超过20MB限制")
        safe_session = _clean_component(session_id, "session")
        safe_name = _clean_component(Path(filename).name, "upload.bin")
        destination = self.uploads_root / safe_session / safe_name
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(content)
        return destination

    def process(self, source_path: Path) -> FileOperationResult:
        suffix = source_path.suffix.lower()
        handler = self._handlers.get(suffix)
        if handler is None:
            supported = "、".join(self.supported_suffixes)
            raise ValueError(f"暂不支持 {suffix or '无扩展名'}；支持：{supported}")
        return handler(source_path)

    def _xlsx_to_markdown(self, source_path: Path) -> FileOperationResult:
        workbook = load_workbook(source_path, read_only=True, data_only=True)
        sections: list[str] = [f"# {source_path.name}"]
        try:
            for worksheet in workbook.worksheets:
                rows = [
                    [_cell_text(value) for value in row[:20]]
                    for row in worksheet.iter_rows(values_only=True)
                ][:100]
                rows = [
                    row
                    for row in rows
                    if any(cell.strip() for cell in row)
                ]
                sections.extend(
                    [
                        "",
                        f"## 工作表：{worksheet.title}",
                        "",
                        *_markdown_table(rows),
                    ]
                )
        finally:
            workbook.close()
        text = "\n".join(sections).strip()
        return FileOperationResult(
            operation_key="xlsx_to_markdown",
            source_name=source_path.name,
            text=text[:MAX_TEXT_CHARS],
        )

    def _docx_to_markdown(self, source_path: Path) -> FileOperationResult:
        document = Document(str(source_path))
        sections: list[str] = [f"# {source_path.name}"]
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (
                paragraph.style.name.lower()
                if paragraph.style is not None and paragraph.style.name
                else ""
            )
            if style_name.startswith("heading"):
                level_text = "".join(character for character in style_name if character.isdigit())
                level = min(max(int(level_text or "2"), 1), 6)
                sections.extend(["", f"{'#' * level} {text}"])
            else:
                sections.extend(["", text])
        for table_index, table in enumerate(document.tables, start=1):
            rows = [
                [_cell_text(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            sections.extend(
                [
                    "",
                    f"## 表格 {table_index}",
                    "",
                    *_markdown_table(rows),
                ]
            )
        text = "\n".join(sections).strip()
        return FileOperationResult(
            operation_key="docx_to_markdown",
            source_name=source_path.name,
            text=text[:MAX_TEXT_CHARS],
        )

    def _markdown_to_pdf(self, source_path: Path) -> FileOperationResult:
        markdown_text = source_path.read_text(encoding="utf-8")
        session_artifacts = self.artifacts_root / source_path.parent.name
        session_artifacts.mkdir(parents=True, exist_ok=True)
        output_path = session_artifacts / f"{source_path.stem}.pdf"
        self._render_markdown_pdf(markdown_text, output_path)
        preview = markdown_text.strip()[:MAX_TEXT_CHARS]
        return FileOperationResult(
            operation_key="markdown_to_pdf",
            source_name=source_path.name,
            text=(
                f"已将 {source_path.name} 转换为 PDF。\n\n"
                f"处理后的文本预览：\n\n{preview}"
            ),
            artifact_path=output_path,
            artifact_mime_type="application/pdf",
        )

    @staticmethod
    def _render_markdown_pdf(markdown_text: str, output_path: Path) -> None:
        font_name = "STSong-Light"
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont(font_name))
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "ChineseBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            alignment=TA_LEFT,
            spaceAfter=4,
        )
        heading = ParagraphStyle(
            "ChineseHeading",
            parent=body,
            fontSize=16,
            leading=22,
            spaceBefore=8,
            spaceAfter=8,
        )
        story: list[Flowable] = []
        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                story.append(Spacer(1, 3 * mm))
                continue
            if line.startswith("#"):
                text = line.lstrip("#").strip()
                story.append(Paragraph(escape(text), heading))
                continue
            if line.startswith(("- ", "* ")):
                line = f"• {line[2:].strip()}"
            story.append(Paragraph(escape(line), body))
        document = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )
        document.build(story)
